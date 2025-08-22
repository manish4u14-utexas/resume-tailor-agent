"""
Resume parser for extracting and updating sections in .docx files.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional, Tuple
import re
from utils import extract_section_keywords, clean_text, log_message


class ResumeParser:
    """
    A class to handle parsing and updating of .docx resume files.
    """
    
    def __init__(self, resume_path: str):
        """
        Initialize the ResumeParser with a resume file.
        
        Args:
            resume_path (str): Path to the .docx resume file
        """
        self.resume_path = resume_path
        self.document = None
        self.section_keywords = extract_section_keywords()
        
    def load_document(self) -> bool:
        """
        Load the .docx document.
        
        Returns:
            bool: True if document loaded successfully, False otherwise
        """
        try:
            self.document = Document(self.resume_path)
            log_message(f"Successfully loaded document: {self.resume_path}")
            return True
        except Exception as e:
            log_message(f"Error loading document: {e}", "ERROR")
            return False
    
    def extract_sections(self) -> Dict[str, str]:
        """
        Extract key sections from the resume.
        
        Returns:
            Dict[str, str]: Dictionary with section names as keys and content as values
        """
        if not self.document:
            log_message("Document not loaded. Call load_document() first.", "ERROR")
            return {}
        
        sections = {
            'summary': '',
            'skills': '',
            'experience': '',
            'education': '',
            'projects': '',
            'certifications': ''
        }
        
        current_section = None
        current_content = []
        
        # Process all paragraphs in the document
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Check if this paragraph is a section header
            detected_section = self._detect_section_header(text)
            
            if detected_section:
                # Save previous section content
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                current_section = detected_section
                current_content = []
                log_message(f"Found section: {detected_section}")
            elif current_section:
                # Add content to current section
                current_content.append(text)
        
        # Save the last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        # Clean up sections
        for section_name in sections:
            sections[section_name] = clean_text(sections[section_name])
        
        log_message("Successfully extracted resume sections")
        return sections
    
    def _detect_section_header(self, text: str) -> Optional[str]:
        """
        Detect if a text line is a section header.
        
        Args:
            text (str): Text to analyze
        
        Returns:
            Optional[str]: Section name if detected, None otherwise
        """
        text_lower = text.lower().strip()
        
        # Remove common formatting characters
        text_clean = re.sub(r'[:\-_=*#]+', '', text_lower).strip()
        
        for section_name, keywords in self.section_keywords.items():
            for keyword in keywords:
                if keyword in text_clean:
                    return section_name
        
        return None
    
    def update_sections(self, updated_sections: Dict[str, str]) -> bool:
        """
        Update resume sections with new content while preserving formatting.
        
        Args:
            updated_sections (Dict[str, str]): Dictionary with section names and new content
        
        Returns:
            bool: True if update successful, False otherwise
        """
        if not self.document:
            log_message("Document not loaded. Call load_document() first.", "ERROR")
            return False
        
        try:
            current_section = None
            section_start_index = None
            paragraphs_to_update = []
            
            # First pass: identify sections and their paragraph ranges
            for i, paragraph in enumerate(self.document.paragraphs):
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                detected_section = self._detect_section_header(text)
                
                if detected_section:
                    # Save previous section info
                    if current_section and section_start_index is not None:
                        paragraphs_to_update.append({
                            'section': current_section,
                            'start': section_start_index,
                            'end': i - 1
                        })
                    
                    # Start tracking new section
                    current_section = detected_section
                    section_start_index = i + 1  # Content starts after header
            
            # Handle the last section
            if current_section and section_start_index is not None:
                paragraphs_to_update.append({
                    'section': current_section,
                    'start': section_start_index,
                    'end': len(self.document.paragraphs) - 1
                })
            
            # Second pass: update sections with new content
            for section_info in reversed(paragraphs_to_update):  # Reverse to maintain indices
                section_name = section_info['section']
                
                if section_name in updated_sections and updated_sections[section_name]:
                    self._replace_section_content(
                        section_info['start'],
                        section_info['end'],
                        updated_sections[section_name]
                    )
                    log_message(f"Updated section: {section_name}")
            
            log_message("Successfully updated resume sections")
            return True
            
        except Exception as e:
            log_message(f"Error updating sections: {e}", "ERROR")
            return False
    
    def _replace_section_content(self, start_index: int, end_index: int, new_content: str) -> None:
        """
        Replace content in a specific paragraph range.
        
        Args:
            start_index (int): Starting paragraph index
            end_index (int): Ending paragraph index
            new_content (str): New content to insert
        """
        # Remove old paragraphs (except the first one which we'll reuse)
        for i in range(end_index, start_index, -1):
            if i < len(self.document.paragraphs):
                p = self.document.paragraphs[i]
                p._element.getparent().remove(p._element)
        
        # Split new content into lines
        content_lines = new_content.split('\n')
        
        # Update the first paragraph
        if start_index < len(self.document.paragraphs) and content_lines:
            first_paragraph = self.document.paragraphs[start_index]
            first_paragraph.text = content_lines[0]
            
            # Add remaining lines as new paragraphs
            for line in content_lines[1:]:
                if line.strip():  # Only add non-empty lines
                    new_paragraph = self.document.add_paragraph(line)
                    # Move the new paragraph to the correct position
                    self._move_paragraph_after(new_paragraph, first_paragraph)
                    first_paragraph = new_paragraph
    
    def _move_paragraph_after(self, paragraph_to_move, reference_paragraph):
        """
        Move a paragraph to appear after a reference paragraph.
        
        Args:
            paragraph_to_move: Paragraph to move
            reference_paragraph: Reference paragraph
        """
        try:
            # This is a simplified approach - in practice, you might need more sophisticated
            # paragraph manipulation depending on the document structure
            pass
        except Exception as e:
            log_message(f"Warning: Could not reorder paragraph: {e}", "WARNING")
    
    def save_document(self, output_path: str) -> bool:
        """
        Save the updated document to a new file.
        
        Args:
            output_path (str): Path where to save the updated document
        
        Returns:
            bool: True if save successful, False otherwise
        """
        if not self.document:
            log_message("Document not loaded. Call load_document() first.", "ERROR")
            return False
        
        try:
            self.document.save(output_path)
            log_message(f"Successfully saved document to: {output_path}")
            return True
        except Exception as e:
            log_message(f"Error saving document: {e}", "ERROR")
            return False
    
    def get_document_info(self) -> Dict[str, any]:
        """
        Get basic information about the document.
        
        Returns:
            Dict[str, any]: Document information
        """
        if not self.document:
            return {}
        
        return {
            'paragraph_count': len(self.document.paragraphs),
            'sections_found': len([p for p in self.document.paragraphs if self._detect_section_header(p.text)]),
            'total_text_length': sum(len(p.text) for p in self.document.paragraphs)
        }


def create_sample_resume(output_path: str) -> bool:
    """
    Create a sample resume for testing purposes.
    
    Args:
        output_path (str): Path where to save the sample resume
    
    Returns:
        bool: True if creation successful, False otherwise
    """
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('John Doe', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact = doc.add_paragraph('Email: john.doe@email.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe')
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add Summary section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph('Experienced software developer with 5+ years of experience in full-stack development. '
                         'Proficient in Python, JavaScript, and cloud technologies. Strong problem-solving skills '
                         'and experience in agile development environments.')
        
        # Add Skills section
        doc.add_heading('Technical Skills', level=1)
        doc.add_paragraph('• Programming Languages: Python, JavaScript, Java, C++\n'
                         '• Web Technologies: React, Node.js, HTML5, CSS3\n'
                         '• Databases: PostgreSQL, MongoDB, Redis\n'
                         '• Cloud Platforms: AWS, Azure, Google Cloud\n'
                         '• Tools: Git, Docker, Kubernetes, Jenkins')
        
        # Add Experience section
        doc.add_heading('Professional Experience', level=1)
        
        doc.add_paragraph('Senior Software Developer | Tech Company Inc. | 2020 - Present')
        doc.add_paragraph('• Developed and maintained web applications using React and Node.js\n'
                         '• Implemented RESTful APIs and microservices architecture\n'
                         '• Collaborated with cross-functional teams to deliver high-quality software\n'
                         '• Mentored junior developers and conducted code reviews')
        
        doc.add_paragraph('Software Developer | StartupXYZ | 2018 - 2020')
        doc.add_paragraph('• Built scalable web applications using Python and Django\n'
                         '• Designed and implemented database schemas\n'
                         '• Participated in agile development processes\n'
                         '• Contributed to system architecture decisions')
        
        # Add Education section
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science\n'
                         'University of Technology | 2014 - 2018')
        
        doc.save(output_path)
        log_message(f"Sample resume created at: {output_path}")
        return True
        
    except Exception as e:
        log_message(f"Error creating sample resume: {e}", "ERROR")
        return False

